from __future__ import annotations

import importlib.util
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
BASE_PATH = ROOT / "report_tools" / "build_course_report.py"
spec = importlib.util.spec_from_file_location("base_report", BASE_PATH)
base = importlib.util.module_from_spec(spec)
assert spec.loader is not None
spec.loader.exec_module(base)

base.SCREENSHOTS = ROOT / "report_assets" / "screenshots-v2"
base.OUTPUT = ROOT / "203214445罗炜皓_基于RAG与Function Calling的校园制度智能问答系统设计与实现_模板版.docx"
base.INK = RGBColor(0, 0, 0)
base.ACCENT = RGBColor(0, 0, 0)
base.MUTED = RGBColor(0, 0, 0)
base.LIGHT_FILL = "FFFFFF"
base.TABLE_HEADER = "FFFFFF"

ORIGINAL_ADD_HEADING = base.add_heading
ORIGINAL_ADD_REFERENCES = base.add_references
ORIGINAL_ADD_PICTURE = base.add_picture


def set_exact_line(paragraph, points: float = 20) -> None:
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    paragraph.paragraph_format.line_spacing = Pt(points)


def configure_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = base.LATIN_FONT
    normal.font.size = Pt(10.5)
    normal._element.rPr.rFonts.set(qn("w:ascii"), base.LATIN_FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), base.LATIN_FONT)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), base.CHINESE_FONT)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.first_line_indent = Pt(21)
    set_exact_line(normal.paragraph_format, 20) if False else None
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    normal.paragraph_format.line_spacing = Pt(20)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)

    heading_settings = {
        "Heading 1": (16, "黑体"),
        "Heading 2": (14, "黑体"),
        "Heading 3": (10.5, base.CHINESE_FONT),
    }
    for name, (size, east_asia) in heading_settings.items():
        style = styles[name]
        style.font.name = base.LATIN_FONT
        style.font.size = Pt(size)
        style.font.bold = name != "Heading 3"
        style.font.color.rgb = RGBColor(0, 0, 0)
        style._element.rPr.rFonts.set(qn("w:ascii"), base.LATIN_FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), base.LATIN_FONT)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
        style.paragraph_format.first_line_indent = Pt(0)
        style.paragraph_format.space_before = Pt(6)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        style.paragraph_format.line_spacing = Pt(20)
        style.paragraph_format.keep_with_next = True

    caption = styles["Caption"]
    caption.font.name = base.LATIN_FONT
    caption.font.size = Pt(10.5)
    caption.font.color.rgb = RGBColor(0, 0, 0)
    caption._element.rPr.rFonts.set(qn("w:ascii"), base.LATIN_FONT)
    caption._element.rPr.rFonts.set(qn("w:hAnsi"), base.LATIN_FONT)
    caption._element.rPr.rFonts.set(qn("w:eastAsia"), base.CHINESE_FONT)
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.first_line_indent = Pt(0)
    caption.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    caption.paragraph_format.line_spacing = Pt(20)


def add_body(doc: Document, text: str, *, bold_lead: str | None = None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Pt(21)
    set_exact_line(p)
    p.paragraph_format.space_after = Pt(0)
    if bold_lead and text.startswith(bold_lead):
        r1 = p.add_run(bold_lead)
        base.set_run_font(r1, size=10.5, bold=True)
        r2 = p.add_run(text[len(bold_lead):])
        base.set_run_font(r2, size=10.5)
    else:
        base.set_run_font(p.add_run(text), size=10.5)
    return p


def add_heading(doc: Document, text: str, level: int = 1):
    mapping = {
        "1 项目背景及研究意义": "1 作品背景",
        "2 相关前沿技术": "2 作品目标与技术选型",
        "5 核心功能设计与实现": "5 开发过程与核心功能实现",
        "7 实验设计与结果分析": "7 测试与优化",
    }
    return ORIGINAL_ADD_HEADING(doc, mapping.get(text, text), level)


def add_picture(doc: Document, path: Path, caption: str, **kwargs) -> None:
    ORIGINAL_ADD_PICTURE(doc, path, caption, **kwargs)
    picture_paragraph = doc.paragraphs[-2]
    picture_paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    picture_paragraph.paragraph_format.line_spacing = 1.0


def add_cover(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.paragraph_format.line_spacing = 1.0
    r = p.add_run()
    r.add_picture(str(ROOT / "report_assets" / "template" / "official-cover-logo.png"), width=Cm(7.8))

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.first_line_indent = Pt(0)
    title.paragraph_format.space_after = Pt(4)
    base.set_run_font(title.add_run("《软件前沿技术》课程期末大作业"), size=18, bold=True)

    year = doc.add_paragraph()
    year.alignment = WD_ALIGN_PARAGRAPH.CENTER
    year.paragraph_format.first_line_indent = Pt(0)
    year.paragraph_format.space_after = Pt(10)
    base.set_run_font(year.add_run("2025—2026 学年第 2 学期"), size=14)

    rows = [
        ("作业题目", "基于 RAG 与 Function Calling 的校园制度智能问答系统设计与实现"),
        ("姓    名", "罗炜皓"), ("学    号", "203214445"),
        ("年    级", "2023 级"), ("班    级", "13902301"),
        ("学    院", "计算机科学与技术学院"), ("专    业", "软件工程"),
        ("指导教师", "赵志强"), ("提交日期", "2026 年 6 月 28 日"),
    ]
    table = doc.add_table(rows=len(rows), cols=2)
    base.set_table_geometry(table, [3.6, 11.0], indent_dxa=0)
    for idx, (label, value) in enumerate(rows):
        for col in range(2):
            tc_pr = table.cell(idx, col)._tc.get_or_add_tcPr()
            borders = OxmlElement("w:tcBorders")
            for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
                tag = OxmlElement(f"w:{edge}")
                tag.set(qn("w:val"), "nil")
                borders.append(tag)
            tc_pr.append(borders)
        left = table.cell(idx, 0).paragraphs[0]
        left.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        left.paragraph_format.first_line_indent = Pt(0)
        set_exact_line(left, 24)
        base.set_run_font(left.add_run(label + "："), size=12, bold=True)
        right = table.cell(idx, 1).paragraphs[0]
        right.alignment = WD_ALIGN_PARAGRAPH.LEFT
        right.paragraph_format.first_line_indent = Pt(0)
        set_exact_line(right, 24)
        base.set_run_font(right.add_run(value), size=12)
    doc.add_page_break()


def add_references(doc: Document, references: list[dict]) -> None:
    for ref in references:
        if ref["id"] == 3:
            ref["url"] = "https://www.nowpublishers.com/article/Details/INR-019"
    ORIGINAL_ADD_REFERENCES(doc, references)
    for p in doc.paragraphs:
        if "链接已于 2026 年 6 月 27 日核验" in p.text:
            for run in p.runs:
                run.text = run.text.replace("2026 年 6 月 27 日", "2026 年 6 月 30 日")


def normalize_document(doc: Document) -> None:
    for p in doc.paragraphs:
        if p._p.xpath(".//w:drawing"):
            continue
        if p.style.name not in {"Heading 1", "Heading 2", "Heading 3"}:
            set_exact_line(p)
        for run in p.runs:
            if p.style.name == "Heading 1":
                base.set_run_font(run, size=16, bold=True)
                run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "黑体")
            elif p.style.name == "Heading 2":
                base.set_run_font(run, size=14, bold=True)
                run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "黑体")
            elif p.style.name == "Heading 3":
                base.set_run_font(run, size=10.5, bold=False)
            elif p.text not in {"摘  要"}:
                base.set_run_font(run, size=10.5)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    set_exact_line(p)


def build() -> Path:
    base.configure_styles = configure_styles
    base.add_body = add_body
    base.add_heading = add_heading
    base.add_picture = add_picture
    base.add_cover = add_cover
    base.add_references = add_references
    base.SCREENSHOTS = ROOT / "report_assets" / "screenshots-v2"

    # Adapt the five fresh browser screenshots to the legacy builder's expected names.
    mapping = {
        "01-login-page.png": "01-login.png",
        "02-chat-home.png": "02-chat-home.png",
        "03-knowledge-base.png": "04-knowledge-base.png",
        "04-pdf-preview.png": "05-pdf-preview.png",
        "05-rag-answer-with-sources.png": "03-rag-answer.png",
    }
    for expected, actual in mapping.items():
        target = base.SCREENSHOTS / expected
        source = base.SCREENSHOTS / actual
        if target == source:
            continue
        if target.exists():
            target.unlink()
        target.symlink_to(source)

    output = base.build()
    doc = Document(output)
    normalize_document(doc)
    doc.save(output)
    return output


if __name__ == "__main__":
    print(build())
