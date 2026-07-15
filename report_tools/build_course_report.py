from __future__ import annotations

import json
import math
import re
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SCREENSHOTS = ROOT / "report_assets" / "screenshots"
FIGURES = ROOT / "report_assets" / "figures"
REFERENCES_PATH = ROOT / "report_assets" / "evidence" / "references.json"
OUTPUT = ROOT / "203214445罗炜皓_基于RAG与Function Calling的校园制度智能问答系统设计与实现.docx"

CHINESE_FONT = "宋体"
LATIN_FONT = "Times New Roman"
SONGTI_PATH = "/System/Library/Fonts/Supplemental/Songti.ttc"

INK = RGBColor(26, 38, 52)
ACCENT = RGBColor(36, 93, 72)
MUTED = RGBColor(92, 103, 112)
LIGHT_FILL = "EEF4F1"
TABLE_HEADER = "DDEBE5"
WHITE = RGBColor(255, 255, 255)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_cm: list[float], indent_dxa: int = 120) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    total_dxa = sum(int(Cm(width).emu / 635) for width in widths_cm)
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total_dxa))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    dxa_widths = [int(Cm(width).emu / 635) for width in widths_cm]
    for width in dxa_widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width_dxa = dxa_widths[min(idx, len(dxa_widths) - 1)]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width_dxa))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Cm(widths_cm[min(idx, len(widths_cm) - 1)])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def set_run_font(run, size: float | None = None, bold: bool | None = None,
                 color: RGBColor | None = None, italic: bool | None = None) -> None:
    run.font.name = LATIN_FONT
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    r_fonts.set(qn("w:ascii"), LATIN_FONT)
    r_fonts.set(qn("w:hAnsi"), LATIN_FONT)
    r_fonts.set(qn("w:eastAsia"), CHINESE_FONT)
    r_fonts.set(qn("w:cs"), LATIN_FONT)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def style_all_runs(paragraph, *, size: float | None = None,
                   bold: bool | None = None, color: RGBColor | None = None) -> None:
    for run in paragraph.runs:
        set_run_font(run, size=size, bold=bold, color=color)


def keep_with_next(paragraph, value: bool = True) -> None:
    paragraph.paragraph_format.keep_with_next = value


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char1, instr_text, fld_char2])
    set_run_font(run, size=9, color=MUTED)


def add_hyperlink(paragraph, text: str, url: str):
    part = paragraph.part
    rel_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), LATIN_FONT)
    r_fonts.set(qn("w:hAnsi"), LATIN_FONT)
    r_fonts.set(qn("w:eastAsia"), CHINESE_FONT)
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    size = OxmlElement("w:sz")
    size.set(qn("w:val"), "20")
    r_pr.extend([r_fonts, color, underline, size])
    text_node = OxmlElement("w:t")
    text_node.text = text
    new_run.extend([r_pr, text_node])
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink


def configure_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = LATIN_FONT
    normal.font.size = Pt(12)
    normal._element.rPr.rFonts.set(qn("w:ascii"), LATIN_FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), LATIN_FONT)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), CHINESE_FONT)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.first_line_indent = Pt(24)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)

    heading_settings = {
        "Heading 1": (16, 14, 8),
        "Heading 2": (14, 10, 5),
        "Heading 3": (12, 8, 4),
    }
    for name, (size, before, after) in heading_settings.items():
        style = styles[name]
        style.font.name = LATIN_FONT
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = INK
        style._element.rPr.rFonts.set(qn("w:ascii"), LATIN_FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), LATIN_FONT)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), CHINESE_FONT)
        style.paragraph_format.first_line_indent = Pt(0)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.25
        style.paragraph_format.keep_with_next = True

    caption = styles["Caption"]
    caption.font.name = LATIN_FONT
    caption.font.size = Pt(10.5)
    caption.font.color.rgb = MUTED
    caption._element.rPr.rFonts.set(qn("w:ascii"), LATIN_FONT)
    caption._element.rPr.rFonts.set(qn("w:hAnsi"), LATIN_FONT)
    caption._element.rPr.rFonts.set(qn("w:eastAsia"), CHINESE_FONT)
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.first_line_indent = Pt(0)
    caption.paragraph_format.space_before = Pt(3)
    caption.paragraph_format.space_after = Pt(8)
    caption.paragraph_format.keep_with_next = False


def configure_sections(doc: Document) -> None:
    for section in doc.sections:
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(2.5)
        section.header_distance = Cm(1.2)
        section.footer_distance = Cm(1.2)
        section.different_first_page_header_footer = True


def setup_header_footer(section) -> None:
    header_p = section.header.paragraphs[0]
    header_p.text = "《软件前沿技术》课程项目报告"
    header_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header_p.paragraph_format.space_after = Pt(0)
    style_all_runs(header_p, size=9, color=MUTED)
    footer_p = section.footer.paragraphs[0]
    add_page_number(footer_p)


def add_body(doc: Document, text: str, *, bold_lead: str | None = None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Pt(24)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(0)
    if bold_lead and text.startswith(bold_lead):
        r1 = p.add_run(bold_lead)
        set_run_font(r1, size=12, bold=True)
        r2 = p.add_run(text[len(bold_lead):])
        set_run_font(r2, size=12)
    else:
        run = p.add_run(text)
        set_run_font(run, size=12)
    return p


def add_lead(doc: Document, label: str, text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [15.5])
    cell = table.cell(0, 0)
    set_cell_shading(cell, LIGHT_FILL)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.line_spacing = 1.3
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    r1 = p.add_run(label + "：")
    set_run_font(r1, size=11, bold=True, color=ACCENT)
    r2 = p.add_run(text)
    set_run_font(r2, size=11, color=INK)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)


def add_heading(doc: Document, text: str, level: int = 1):
    p = doc.add_paragraph(text, style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    style_all_runs(p, bold=True, color=INK)
    return p


def add_table(doc: Document, headers: list[str], rows: list[list[str]],
              widths: list[float], alignments: list[int] | None = None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    set_repeat_table_header(table.rows[0])
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, TABLE_HEADER)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Pt(0)
        p.paragraph_format.line_spacing = 1.15
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(header)
        set_run_font(run, size=10.5, bold=True, color=INK)
    for row_values in rows:
        row = table.add_row()
        for idx, value in enumerate(row_values):
            cell = row.cells[idx]
            p = cell.paragraphs[0]
            p.alignment = (alignments[idx] if alignments else WD_ALIGN_PARAGRAPH.LEFT)
            p.paragraph_format.first_line_indent = Pt(0)
            p.paragraph_format.line_spacing = 1.15
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(str(value))
            set_run_font(run, size=10.5)
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(2)
    return table


def add_picture(doc: Document, path: Path, caption: str, *, max_width_cm=15.5,
                max_height_cm=11.5, alt_text: str | None = None) -> None:
    if not path.exists():
        raise FileNotFoundError(path)
    with Image.open(path) as image:
        width, height = image.size
    width_cm = max_width_cm
    height_cm = width_cm * height / width
    if height_cm > max_height_cm:
        height_cm = max_height_cm
        width_cm = height_cm * width / height
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.keep_with_next = True
    run = p.add_run()
    inline = run.add_picture(str(path), width=Cm(width_cm), height=Cm(height_cm))
    doc_pr = inline._inline.docPr
    doc_pr.set("descr", alt_text or caption)
    cap = doc.add_paragraph(caption, style="Caption")
    style_all_runs(cap, size=10.5, color=MUTED)


def font(size: int, index: int = 0):
    try:
        return ImageFont.truetype(SONGTI_PATH, size=size, index=index)
    except Exception:
        return ImageFont.load_default()


def rounded_box(draw, xy, fill, outline, radius=18, width=3):
    draw.rounded_rectangle(xy, radius=radius, fill=fill, outline=outline, width=width)


def arrow(draw, start, end, color=(70, 104, 90), width=5):
    draw.line([start, end], fill=color, width=width)
    angle = math.atan2(end[1] - start[1], end[0] - start[0])
    length = 15
    for delta in (2.55, -2.55):
        point = (
            end[0] + length * math.cos(angle + delta),
            end[1] + length * math.sin(angle + delta),
        )
        draw.line([end, point], fill=color, width=width)


def create_architecture_figure() -> Path:
    FIGURES.mkdir(parents=True, exist_ok=True)
    path = FIGURES / "system-architecture.png"
    canvas = Image.new("RGB", (1800, 980), "#F8FAF9")
    draw = ImageDraw.Draw(canvas)
    title_font = font(46)
    box_title = font(31)
    box_text = font(24)
    draw.text((900, 45), "CQUPT-RAG 系统总体架构", font=title_font, fill="#1A2634", anchor="mm")

    boxes = [
        (80, 190, 370, 760, "前端交互层", ["React + TypeScript", "登录与会话管理", "WebSocket 流式对话", "来源引用与 PDF 预览"]),
        (500, 140, 820, 810, "应用服务层", ["FastAPI 路由", "JWT 身份认证", "Agent 双轮调用", "长短期记忆编排"]),
        (950, 110, 1320, 845, "RAG 检索层", ["查询改写", "向量召回", "BM25 关键词检索", "RRF 融合与重排", "上下文与来源组织"]),
        (1450, 190, 1730, 760, "数据与模型层", ["GLM-4.7-Flash", "智谱 embedding-2", "Chroma 向量库", "MySQL 会话数据", "校园制度文档"]),
    ]
    fills = ["#EAF3EF", "#E4EFE9", "#DCEBE4", "#F0F4F2"]
    outlines = ["#4C8B70", "#37755D", "#2E6B53", "#71847B"]
    for idx, (x1, y1, x2, y2, title, items) in enumerate(boxes):
        rounded_box(draw, (x1, y1, x2, y2), fills[idx], outlines[idx], radius=24, width=4)
        draw.text(((x1 + x2) / 2, y1 + 58), title, font=box_title, fill="#1C3D30", anchor="mm")
        y = y1 + 150
        for item in items:
            draw.ellipse((x1 + 35, y - 8, x1 + 49, y + 6), fill=outlines[idx])
            draw.text((x1 + 68, y), item, font=box_text, fill="#263730", anchor="lm")
            y += 88
    arrow(draw, (370, 480), (500, 480))
    arrow(draw, (820, 480), (950, 480))
    arrow(draw, (1320, 480), (1450, 480))
    draw.text((900, 920), "数据流：用户问题 → 检索增强 → 模型生成 → 来源可追溯回答", font=box_text, fill="#4A5B53", anchor="mm")
    canvas.save(path, quality=95)
    return path


def create_evaluation_figure() -> Path:
    FIGURES.mkdir(parents=True, exist_ok=True)
    path = FIGURES / "evaluation-comparison.png"
    canvas = Image.new("RGB", (1600, 900), "#FFFFFF")
    draw = ImageDraw.Draw(canvas)
    title_font = font(44)
    label_font = font(28)
    value_font = font(30)
    note_font = font(23)
    draw.text((800, 55), "四种检索策略准确率对比", font=title_font, fill="#1A2634", anchor="mm")
    values = [("baseline", 41.7), ("rerank", 58.3), ("rewrite", 66.7), ("hybrid", 41.7)]
    chart_left, chart_top, chart_right, chart_bottom = 180, 160, 1480, 710
    for tick in range(0, 81, 20):
        y = chart_bottom - (tick / 80) * (chart_bottom - chart_top)
        draw.line((chart_left, y, chart_right, y), fill="#DDE3E0", width=2)
        draw.text((chart_left - 35, y), f"{tick}%", font=note_font, fill="#5A6962", anchor="rm")
    bar_width = 190
    centers = [360, 650, 940, 1230]
    colors = ["#91A69C", "#5E927B", "#2F7659", "#A7B2AD"]
    for (name, value), center, color in zip(values, centers, colors):
        y = chart_bottom - (value / 80) * (chart_bottom - chart_top)
        draw.rounded_rectangle((center - bar_width / 2, y, center + bar_width / 2, chart_bottom), radius=16, fill=color)
        draw.text((center, y - 28), f"{value:.1f}%", font=value_font, fill="#1A2634", anchor="ms")
        draw.text((center, chart_bottom + 48), name, font=label_font, fill="#263730", anchor="mt")
    draw.text((800, 825), "测试集：12 条结构化校园制度问答；rewrite 策略在本次实验中表现最佳", font=note_font, fill="#4A5B53", anchor="mm")
    canvas.save(path, quality=95)
    return path


def add_cover(doc: Document) -> None:
    for _ in range(3):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(12)
    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.first_line_indent = Pt(0)
    kicker.paragraph_format.space_after = Pt(24)
    r = kicker.add_run("《软件前沿技术》课程期末大作业")
    set_run_font(r, size=16, bold=True, color=ACCENT)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.first_line_indent = Pt(0)
    title.paragraph_format.line_spacing = 1.35
    title.paragraph_format.space_after = Pt(12)
    r = title.add_run("基于 RAG 与 Function Calling 的\n校园制度智能问答系统设计与实现")
    set_run_font(r, size=24, bold=True, color=INK)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.first_line_indent = Pt(0)
    subtitle.paragraph_format.space_after = Pt(58)
    r = subtitle.add_run("CQUPT-RAG Assistant")
    set_run_font(r, size=14, color=MUTED)

    rows = [
        ("学    号", "203214445"),
        ("姓    名", "罗炜皓"),
        ("班    级", "13902301"),
        ("学    院", "计算机科学与技术学院"),
        ("专    业", "软件工程"),
        ("指导教师", "赵志强"),
        ("提交日期", "2026 年 6 月 28 日"),
    ]
    table = doc.add_table(rows=len(rows), cols=2)
    set_table_geometry(table, [4.2, 8.0], indent_dxa=0)
    for idx, (label, value) in enumerate(rows):
        for col in range(2):
            cell = table.cell(idx, col)
            tc_pr = cell._tc.get_or_add_tcPr()
            borders = OxmlElement("w:tcBorders")
            for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
                tag = OxmlElement(f"w:{edge}")
                tag.set(qn("w:val"), "nil")
                borders.append(tag)
            tc_pr.append(borders)
        p1 = table.cell(idx, 0).paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p1.paragraph_format.first_line_indent = Pt(0)
        r1 = p1.add_run(label + "：")
        set_run_font(r1, size=13, bold=True, color=INK)
        p2 = table.cell(idx, 1).paragraphs[0]
        p2.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p2.paragraph_format.first_line_indent = Pt(0)
        r2 = p2.add_run(value)
        set_run_font(r2, size=13, color=INK)
    doc.add_page_break()


def add_abstract(doc: Document) -> None:
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.first_line_indent = Pt(0)
    title.paragraph_format.space_after = Pt(12)
    r = title.add_run("摘  要")
    set_run_font(r, size=16, bold=True, color=INK)
    add_body(doc, "随着高校制度文件数量不断增加，学生在查询奖学金、学籍异动、考试管理和违纪处分等规定时，常面临文件分散、关键词不确定、人工检索耗时以及答案缺少出处等问题。针对这些问题，本文设计并实现了一套基于检索增强生成（Retrieval-Augmented Generation，RAG）与 Function Calling 的校园制度智能问答系统。系统以重庆邮电大学学生手册和校级制度文件为知识源，通过文档解析、文本切分、智谱 embedding-2 向量化和 Chroma 持久化索引构建校园知识库；在查询阶段实现向量检索、领域关键词重排、查询改写以及 BM25 与向量检索的 RRF 融合，并将检索片段连同页码和文件名注入 GLM-4.7-Flash 生成流程。系统采用 FastAPI、React、MySQL 和 WebSocket 构建完整 Web 应用，支持用户认证、会话管理、流式回答、长短期记忆、来源引用、知识库浏览和 PDF 原页预览。")
    add_body(doc, "为评价检索策略的有效性，项目构建了包含事实、规则、多条件和拒答四类问题的 12 条测试集。实验结果表明，纯向量检索准确率为 41.7%，加入领域关键词重排后达到 58.3%，查询改写策略达到 66.7%；混合检索在当前小规模语料和参数设置下未取得提升。结果说明查询表达扩展能够改善校园制度问答的召回覆盖率，同时也暴露出中文分词、融合参数和评测规模仍需优化。该系统展示了大语言模型、语义检索与智能体工具调用在校园信息服务中的可落地性。")
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.space_before = Pt(10)
    r1 = p.add_run("关键词：")
    set_run_font(r1, size=12, bold=True)
    r2 = p.add_run("检索增强生成；大语言模型；混合检索；Function Calling；校园智能问答")
    set_run_font(r2, size=12)
    add_lead(doc, "项目结论", "该项目符合《软件前沿技术》课程中人工智能与 AIGC 方向要求，既包含前沿技术原理，也形成了可在 PC 端现场展示的完整作品。")
    doc.add_page_break()


def add_main_content(doc: Document, architecture_path: Path, evaluation_path: Path) -> None:
    add_heading(doc, "1 项目背景及研究意义", 1)
    add_heading(doc, "1.1 项目背景", 2)
    add_body(doc, "高校学生管理制度具有条款多、版本更新快和跨文件关联强等特点。传统查询方式通常依赖网页搜索、人工翻阅 PDF 或向辅导员咨询。对于“哪些处分可以获得补考机会”“国家奖学金能否与其他奖学金兼得”等问题，用户不仅需要定位关键词，还要同时理解适用对象、条件、例外和政策依据。普通关键词搜索容易遗漏同义表达，而通用大语言模型虽然具备自然语言理解能力，却可能依据训练语料生成与学校现行制度不一致的答案。")
    add_body(doc, "RAG 将外部知识检索与生成模型结合，使模型能够在回答前读取指定知识库中的相关片段。Lewis 等提出的 RAG 框架将参数化语言模型与非参数化检索记忆结合，为知识密集型任务提供了可更新知识与来源依据的技术路线[1]。因此，本项目选择校园制度这一边界清晰、对准确性和可追溯性要求较高的场景，探索大模型应用从“能回答”走向“有依据地回答”。")

    add_heading(doc, "1.2 研究意义", 2)
    add_body(doc, "从应用角度看，系统把分散的学生手册和专项制度统一组织为可检索知识库，用户可以使用自然语言提问，并通过引用页码直接回到原文核查。它能够降低制度查询门槛，提高信息获取效率。从工程角度看，项目完整覆盖文档处理、向量化、检索、生成、实时通信、身份认证和前端交互，体现了前沿人工智能技术与传统软件工程方法的结合。")
    add_body(doc, "从研究角度看，项目并未把大模型输出直接等同于正确答案，而是设计四种检索策略和结构化测试集进行对比。实验既记录准确率，也记录响应时延，用以分析检索质量与系统成本之间的权衡。这种以可复核实验检验系统方案的过程，比单纯展示界面更能体现课程项目的技术深度。")

    add_heading(doc, "2 相关前沿技术", 1)
    add_heading(doc, "2.1 大语言模型与 AIGC", 2)
    add_body(doc, "大语言模型通过大规模预训练获得文本理解、生成和指令遵循能力，是当前 AIGC 应用的核心基础。本项目调用智谱 GLM-4.7-Flash 完成查询改写、意图判断、工具选择、历史摘要和最终回答。需要强调的是，项目没有重新训练基础模型，其创新重点在于如何为预训练模型提供可靠的校园知识、约束回答边界并形成完整的软件系统。")

    add_heading(doc, "2.2 检索增强生成", 2)
    add_body(doc, "RAG 的基本流程是先根据用户问题从外部知识库中检索相关文档，再把检索结果作为上下文交给生成模型。与仅依赖模型参数的方式相比，它具有知识可更新、答案可溯源和领域适配成本较低等优势[1]。本项目在提示词中明确要求模型只能依据政策资料作答；资料缺少明确答案时应返回“无法确定”，从流程上降低无依据补充的风险。")

    add_heading(doc, "2.3 文本向量与向量数据库", 2)
    add_body(doc, "文本向量把语句映射到高维数值空间，使语义相近的文本具有较高相似度。Sentence-BERT 证明了句向量在语义相似度检索中的有效性[2]。本项目实际使用智谱 embedding-2 将制度文本块和用户查询编码为向量，并利用 Chroma 建立持久化集合。Chroma 的查询接口支持依据查询向量执行近邻搜索，并返回文本、元数据与相似度相关结果[5]。")

    add_heading(doc, "2.4 BM25 与 RRF 融合", 2)
    add_body(doc, "向量检索擅长处理语义相似表达，但对精确政策名、金额和专有术语不一定稳定；BM25 根据词频和逆文档频率评估相关性，对关键词匹配更敏感。Robertson 和 Zaragoza 对 BM25 的概率相关框架进行了系统总结[3]。为同时利用两种检索信号，项目采用 Reciprocal Rank Fusion（RRF）按排名倒数累计得分。RRF 不要求不同检索器的原始分数处于同一量纲，适合融合词法检索与向量检索结果[4]。")

    add_heading(doc, "2.5 Function Calling 与 WebSocket", 2)
    add_body(doc, "Function Calling 允许模型从预先声明的函数中选择工具并生成结构化参数，程序执行函数后再把结果交还模型组织最终回答。智谱官方文档给出了 tools、tool_calls 和 tool 消息的完整交互方式[8]。项目据此实现天气和课表两个演示工具。实时对话部分使用 WebSocket 保持前后端双向连接，FastAPI 官方文档说明了服务端接受连接、接收消息和持续发送结果的实现模式[7]。")

    add_heading(doc, "3 系统需求分析", 1)
    add_heading(doc, "3.1 用户需求", 2)
    add_body(doc, "系统主要面向需要快速查询校内制度的学生。用户希望以自然语言描述问题，而不必准确知道文件名称或条款编号；得到答案后还应能够查看来源文件、页码和原文，以判断答案是否适用于自己的情况。系统同时需要保存历史会话，使连续提问能够保持上下文。")
    add_table(doc, ["需求类别", "具体要求", "实现方式"], [
        ["身份与会话", "注册、登录、新建会话、历史记录与搜索", "JWT、MySQL、会话 REST 接口"],
        ["智能问答", "自然语言提问、流式显示、连续对话", "GLM-4.7-Flash、WebSocket、记忆系统"],
        ["知识检索", "覆盖语义表达和精确政策关键词", "向量检索、关键词重排、BM25、RRF"],
        ["可信依据", "显示来源文件、页码和原文片段", "稳定元数据、引用卡片、PDF 页预览"],
        ["资料管理", "浏览当前挂载的制度资料", "知识库列表与文件状态检查"],
    ], [3.0, 6.3, 6.2], [WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT])

    add_heading(doc, "3.2 非功能需求", 2)
    add_body(doc, "准确性方面，答案必须优先依据当前知识库，缺少材料时不应凭常识补全；可用性方面，页面应提供问题示例、生成状态和错误反馈；性能方面，耗时较长的模型调用通过流式输出减少等待感；安全方面，认证接口使用密码哈希和 JWT，受保护接口校验当前用户；可维护性方面，前后端分层、检索策略可切换，文档元数据使用稳定字段。")

    add_heading(doc, "4 系统总体设计", 1)
    add_heading(doc, "4.1 总体架构", 2)
    add_body(doc, "系统采用前后端分离架构。前端负责身份输入、会话导航、消息渲染和来源交互；FastAPI 负责认证、会话、文档和 WebSocket 路由；RAG 层负责查询改写、召回、融合、上下文组织和生成；数据与模型层包含 GLM、Embedding 接口、Chroma、MySQL 及本地制度文件。各层通过明确接口连接，便于独立替换模型或检索策略。")
    add_picture(doc, architecture_path, "图 4-1 系统总体架构", max_height_cm=8.3, alt_text="CQUPT-RAG 前端、应用服务、RAG 检索及数据模型四层架构图")

    add_heading(doc, "4.2 问答数据流程", 2)
    add_body(doc, "用户发送问题后，后端首先读取最近会话消息和长期摘要，随后根据当前策略构造检索请求。查询改写策略把一个问题扩展成 2—4 个角度不同的子查询；每个子查询分别召回候选文本块，并按内容去重。系统把最终 3 条文本块格式化为带有文档名和页码的上下文，再与用户问题和记忆信息共同组成提示词。模型输出通过 WebSocket 按增量发送，生成结束时单独发送来源列表，前端据此渲染引用卡片。")

    add_heading(doc, "4.3 数据组织", 2)
    add_body(doc, "最近一次索引元数据表明，知识库包含 4 份制度文档、201 页和 329 个文本块。每个文本块除正文外，还保存 document_id、document_name、document_type、topic、authority_level、source 和 page 等字段。稳定元数据使回答来源能够从向量检索结果一直传递到前端，而不是在生成完成后通过文本猜测来源。")

    add_heading(doc, "5 核心功能设计与实现", 1)
    add_heading(doc, "5.1 文档解析与知识库构建", 2)
    add_body(doc, "系统支持 PDF 与 DOCX 两类制度文件。PDF 优先使用 PyMuPDFLoader 解析，失败时回退到 PyPDFLoader；DOCX 则直接读取 OOXML 中的段落文本。所有文档统一注入稳定元数据，再使用 RecursiveCharacterTextSplitter 切分为长度 600、重叠 80 的文本块。重叠区域用于保留跨边界上下文，避免条款条件与结论被完全拆开。")
    add_body(doc, "向量化阶段通过 BatchedZhipuAIEmbeddings 将文本按每批 16 条提交给 embedding-2，降低单次请求超过接口限制的风险。生成的向量和文档元数据写入 Chroma 持久化目录。应用启动时优先检测现有集合；当集合存在且包含数据时直接加载，否则重新解析文件并构建索引。这种方式兼顾首次建库和日常快速启动。")

    add_heading(doc, "5.2 多策略检索", 2)
    add_table(doc, ["策略", "处理流程", "主要目的"], [
        ["baseline", "原问题 → 向量检索 → Top 3", "作为基础对照"],
        ["rerank", "向量召回 8 条 → 领域关键词重排 → Top 3", "增强政策术语匹配"],
        ["rewrite", "改写 2—4 个子查询 → 分别向量召回 → 去重重排", "扩大语义召回覆盖"],
        ["hybrid", "子查询 → 向量与 BM25 双路召回 → RRF 融合", "结合语义和关键词信号"],
    ], [2.6, 8.6, 4.3], [WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT])
    add_body(doc, "关键词重排模块内置奖学金、学籍、处分、补考和申诉等校园领域词表。当问题命中专有术语时，包含相同术语的候选文档获得更高排序权重。查询改写由 GLM-4.7-Flash 完成，并强制保留原问题作为第一条子查询。混合检索使用 jieba 对中文文本分词，BM25 与向量结果以 RRF 公式 score(d)=Σ1/(k+rank_i) 融合，代码中 k 取 60。")

    add_heading(doc, "5.3 受约束生成与来源引用", 2)
    add_body(doc, "生成提示词明确限定模型身份和回答边界：只能根据提供的政策资料作答；资料没有明确答案时回答“无法确定”；奖学金问题应按结论、条件、互斥限制和依据文件组织。检索上下文采用“【文档名｜第 n 页】”格式，使模型和后续来源解析模块都能识别出处。生成完成后，后端从知识上下文中提取文档名、页码和片段，形成结构化 sources 数组。")
    add_body(doc, "前端来源卡片支持悬停缩略图和点击原页。当用户选择某一引用时，受保护的文档接口根据白名单 document_id 定位文件，并使用 PyMuPDF 按需渲染指定页。页面缓存减少重复渲染成本，页码越界或文件不存在时返回明确错误。")

    add_heading(doc, "5.4 Agent 工具调用与记忆", 2)
    add_body(doc, "Agent 链路采用两轮模型调用。第一轮把工具 schema 与对话消息一并发送给模型；若 finish_reason 为 tool_calls，后端解析 JSON 参数并执行对应函数，再将函数结果以 tool 角色消息回填；第二轮模型根据工具结果生成自然语言回答。当前天气和课表工具使用本地模拟数据，作用是验证 Function Calling 流程，不宣称已经接入真实外部系统。")
    add_body(doc, "记忆系统分为短期与长期两路。短期记忆从 MySQL 读取最近消息，保持当前会话连贯性；长期记忆按消息窗口生成摘要，写入独立 Chroma 集合，并在后续问题中通过相似度检索相关摘要。两路记忆与知识库上下文合并后进入系统提示词，使系统既能依据制度回答，也能参考用户此前对话。")

    add_heading(doc, "5.5 用户认证与前后端交互", 2)
    add_body(doc, "认证模块支持手机号检查、注册、登录和当前用户查询。密码经过哈希后保存，登录成功签发 JWT。前端 Axios 拦截器自动附加 Authorization 头，遇到 401 响应时清理本地登录状态。聊天页面使用 Zustand 管理认证、会话和消息状态，React Markdown 渲染结构化回答，WebSocket 客户端处理检索步骤、文本增量、来源和结束事件。")

    doc.add_page_break()
    add_heading(doc, "6 系统功能展示", 1)
    add_heading(doc, "6.1 登录界面", 2)
    add_body(doc, "系统提供统一登录入口。手机号失焦后，前端调用检查接口判断账号是否存在，并自动切换“登录”或“注册并登录”流程；密码输入支持显隐切换。界面在 PC 端采用左右分区和暗色视觉设计，便于现场演示。")
    add_picture(doc, SCREENSHOTS / "01-login-page.png", "图 6-1 系统登录界面", max_height_cm=8.2, alt_text="CQUPT-RAG 手机号和密码登录页面")

    add_heading(doc, "6.2 智能问答主界面", 2)
    add_body(doc, "登录后进入问答主页。页面左侧提供新建对话、搜索聊天、知识库和历史会话入口；主区域给出奖学金、补考重修、违纪处分和学籍异动等示例问题。输入框通过 WebSocket 发送问题，生成期间显示检索状态，结束后保存会话。")
    add_picture(doc, SCREENSHOTS / "02-chat-home.png", "图 6-2 智能问答主界面", max_height_cm=8.2, alt_text="CQUPT-RAG 智能问答首页和问题示例")

    add_heading(doc, "6.3 知识库资料浏览", 2)
    add_body(doc, "知识库弹窗实时读取后端白名单中的制度文件，展示名称、制度类型、页数、文件大小和预览状态。当前系统实际挂载 4 份资料，其中 3 份 PDF 可直接预览，1 份 DOCX 暂不支持页图预览。该界面让使用者清楚知道回答依据来自哪些材料。")
    add_picture(doc, SCREENSHOTS / "03-knowledge-base.png", "图 6-3 知识库资料列表", max_height_cm=8.0, alt_text="知识库中四份校园制度资料的列表")

    add_heading(doc, "6.4 PDF 原文预览", 2)
    add_body(doc, "点击可预览资料后，系统打开 PDF 页图弹窗，显示文档名称、当前页码、总页数以及上一页和下一页控件。图 6-4 展示的是《学生手册（教育管理篇）2025版》第 1 页，页面内容由后端从真实 PDF 按需渲染，并非前端静态设计图。")
    add_picture(doc, SCREENSHOTS / "04-pdf-preview.png", "图 6-4 PDF 原文件页面预览", max_height_cm=8.4, alt_text="学生手册真实 PDF 第一页预览")

    add_heading(doc, "6.5 带来源的 RAG 回答", 2)
    add_body(doc, "图 6-5 展示系统对“国家奖学金的奖励标准是多少，申请对象有哪些要求”的实际回答。回答完成后，页面列出《学生手册（教育管理篇）2025版》第 55、82 和 83 页三个来源入口。用户可以从自然语言结论继续进入原文核查，形成“检索—回答—证据”的闭环。")
    add_picture(doc, SCREENSHOTS / "05-rag-answer-with-sources.png", "图 6-5 带页码来源的 RAG 回答", max_height_cm=8.4, alt_text="国家奖学金问题回答及三个学生手册页码来源")

    add_heading(doc, "7 实验设计与结果分析", 1)
    add_heading(doc, "7.1 测试集与评价方法", 2)
    add_body(doc, "测试集共 12 条问题，由 5 条事实类、4 条规则类、2 条多条件类和 1 条拒答类构成。每条用例给出预期答案、必须包含词和禁止包含词。系统分别使用四种检索策略生成回答，再由统一评价流程判断答案是否满足关键事实和拒答要求。准确率定义为通过用例数除以总用例数；平均时延记录从检索到回答评价的完整耗时。由于测试集规模较小，结果只用于项目内部方案比较，不代表面向全部校园问题的总体准确率。")

    add_heading(doc, "7.2 实验结果", 2)
    add_table(doc, ["检索策略", "通过数", "准确率", "平均时延/s", "事实类通过率"], [
        ["baseline", "5/12", "41.7%", "8.16", "3/5（60%）"],
        ["rerank", "7/12", "58.3%", "12.23", "5/5（100%）"],
        ["rewrite", "8/12", "66.7%", "25.47", "5/5（100%）"],
        ["hybrid", "5/12", "41.7%", "31.01", "3/5（60%）"],
    ], [3.3, 2.7, 2.8, 3.2, 3.5], [WD_ALIGN_PARAGRAPH.CENTER] * 5)
    add_picture(doc, evaluation_path, "图 7-1 四种检索策略准确率对比", max_height_cm=8.0, alt_text="baseline、rerank、rewrite 和 hybrid 准确率柱状图")

    add_heading(doc, "7.3 结果讨论", 2)
    add_body(doc, "实验表明，领域关键词重排将准确率从 41.7% 提高到 58.3%，说明奖学金、处分、重修等高辨识度术语能够帮助正确片段进入最终上下文。查询改写进一步达到 66.7%，事实类问题 5 条全部通过，说明从多个角度表达原问题能够缓解用户表述与制度原文不一致的问题。其代价是平均时延上升到 25.47 秒，因为需要额外模型调用和多次向量检索。")
    add_body(doc, "混合检索在本次实验中仍为 41.7%，且平均时延最高。该结果并不意味着 BM25 或 RRF 无效，而是提示当前实现仍存在适配问题：语料只有 329 个文本块，中文分词与政策专有词典不足；BM25 和向量检索候选集合高度重合，互补性有限；RRF 固定 k=60、候选数和最终 Top 3 尚未调参；多级查询改写与融合也增加了延迟。后续应扩充测试集，记录检索级 Recall@K、MRR 等指标，并对分词词典、k 值和候选数量进行网格实验。")
    add_lead(doc, "指标说明", "仓库中的旧版 test_report.md 仅记录 18/20（90%）且缺少用例明细，无法与本次 12 条实验直接比较。因此正式报告只采用 experiment_report.json 中可复核的四策略结果。")

    add_heading(doc, "8 项目创新点、不足与改进方向", 1)
    add_heading(doc, "8.1 项目创新点", 2)
    add_body(doc, "第一，系统不是单纯调用大模型，而是形成包含文档解析、向量索引、多策略检索、受约束生成、来源解析和原页预览的完整 RAG 证据链。第二，检索层可以在四种策略之间切换，并以结构化测试集比较效果，使优化过程具备量化依据。第三，系统把 Function Calling、长短期记忆、会话管理和 PDF 预览整合到同一 Web 应用，体现了前沿模型能力与软件工程架构的融合。第四，来源引用保存了文档名和页码，用户能够从答案直接回到原文，增强了校园制度问答的可核查性。")

    add_heading(doc, "8.2 当前不足", 2)
    add_body(doc, "当前测试集只有 12 条，覆盖范围和统计可信度有限；LLM-as-Judge 或关键词规则可能产生评价偏差。混合检索尚未完成系统调参，时延较高。知识库文件仍依赖本地配置，缺少面向管理员的上传、版本审核和索引状态页面。天气与课表工具使用模拟数据，没有连接真实服务。虽然提示词要求模型依据资料回答，但实际回答仍可能把不同条款组合错误，因此关键制度结论必须由用户查看原文确认。")

    add_heading(doc, "8.3 改进方向", 2)
    add_body(doc, "后续可从四方面改进：扩大由人工审核的测试集，并分别评估检索与生成；加入交叉编码器重排或更适合中文政策的 Embedding 模型；对 BM25 词典、RRF 参数和召回数量进行系统实验；建设管理员端文档版本管理与增量索引机制。若接入真实教务和天气 API，还需增加权限控制、数据脱敏、调用审计和失败降级。")

    add_heading(doc, "9 总结", 1)
    add_body(doc, "本文完成了基于 RAG 与 Function Calling 的校园制度智能问答系统设计与实现。系统以 4 份校园制度文件为知识源，构建包含 201 页、329 个文本块的 Chroma 向量知识库，实现向量检索、领域重排、查询改写和 BM25+RRF 混合检索；结合 GLM-4.7-Flash、WebSocket、JWT、MySQL 和 React，形成可登录、可连续对话、可浏览资料并可预览 PDF 原文的 PC 端作品。")
    add_body(doc, "对比实验显示，查询改写策略在当前 12 条测试集上取得 66.7% 的最高准确率，较 baseline 提升 25 个百分点；与此同时，混合检索没有获得预期提升，说明检索系统优化必须依赖可复核实验，而不能仅凭技术叠加判断效果。总体而言，该项目满足《软件前沿技术》课程对人工智能和 AIGC 方向作品的要求，也为校园制度信息服务提供了一个可继续扩展的工程基础。")


def add_references(doc: Document, references: list[dict]) -> None:
    add_heading(doc, "参考文献", 1)
    intro = doc.add_paragraph()
    intro.paragraph_format.first_line_indent = Pt(0)
    intro.paragraph_format.space_after = Pt(6)
    r = intro.add_run("以下文献均为可公开检索的论文出版页或官方技术文档，链接已于 2026 年 6 月 27 日核验。")
    set_run_font(r, size=10.5, color=MUTED)
    for ref in references:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.first_line_indent = Pt(-18)
        p.paragraph_format.left_indent = Pt(18)
        p.paragraph_format.line_spacing = 1.25
        p.paragraph_format.space_after = Pt(5)
        citation = f"[{ref['id']}] {ref['authors']} {ref['title']}["
        citation += "EB/OL" if ref["type"] == "documentation" else ("C" if ref["type"] == "conference" else "J")
        citation += f"]. {ref['venue']}, {ref['year']}. "
        r = p.add_run(citation)
        set_run_font(r, size=10.5)
        add_hyperlink(p, ref["url"], ref["url"])


def add_appendix(doc: Document) -> None:
    doc.add_page_break()
    add_heading(doc, "附录 A 现场演示建议", 1)
    add_body(doc, "现场展示时建议按固定路径操作，以便在较短时间内体现完整技术闭环。首先打开登录界面并说明手机号注册、JWT 鉴权和会话隔离；随后进入知识库列表，展示 4 份制度资料及页数；接着提出“国家奖学金的奖励标准是多少，请给出依据”等问题，观察检索提示和流式输出；回答结束后指出来源页码并点击打开 PDF 原文；最后展示历史会话和搜索功能。整个演示重点应放在“答案能够回到原文核查”，而不是只介绍界面效果。")
    add_table(doc, ["演示环节", "建议操作", "需要说明的技术点"], [
        ["登录", "输入演示账号并进入系统", "密码哈希、JWT、受保护接口"],
        ["知识库", "打开资料列表", "4 份文档、201 页、329 个文本块"],
        ["问答", "提问国家奖学金或补考规定", "查询改写、向量检索、流式生成"],
        ["来源", "点击页码引用", "稳定元数据、PDF 按需渲染"],
        ["会话", "展示历史记录与搜索", "MySQL 会话数据、长短期记忆"],
    ], [2.5, 6.4, 6.6], [WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT])
    add_lead(doc, "演示提醒", "天气和课表工具当前使用模拟数据，现场应如实说明其用途是验证 Function Calling 链路，不要表述为已经接入真实教务系统。")


def set_core_properties(doc: Document) -> None:
    props = doc.core_properties
    props.title = "基于 RAG 与 Function Calling 的校园制度智能问答系统设计与实现"
    props.subject = "《软件前沿技术》课程期末大作业"
    props.author = "罗炜皓"
    props.keywords = "RAG, Function Calling, 校园智能问答, AIGC"
    props.comments = "学号 203214445；班级 13902301；指导教师 赵志强"


def build() -> Path:
    references = json.loads(REFERENCES_PATH.read_text(encoding="utf-8"))
    architecture = create_architecture_figure()
    evaluation = create_evaluation_figure()
    doc = Document()
    configure_styles(doc)
    configure_sections(doc)
    setup_header_footer(doc.sections[0])
    set_core_properties(doc)

    add_cover(doc)
    add_abstract(doc)
    add_main_content(doc, architecture, evaluation)
    add_references(doc, references)
    add_appendix(doc)

    settings = doc.settings._element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")

    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    output = build()
    print(output)

